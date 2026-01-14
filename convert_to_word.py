#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script para convertir el reporte de estructura de Markdown a Word
"""

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Instalando python-docx...")
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

import re
from pathlib import Path

def convert_md_to_docx(md_file, docx_file):
    """Convierte un archivo Markdown a Word"""
    
    # Leer el archivo markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Crear documento Word
    doc = Document()
    
    # Procesar línea por línea
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Bloques de código
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                # Fin del bloque de código
                in_code_block = False
                if code_lines:
                    p = doc.add_paragraph('\n'.join(code_lines))
                    p.style = 'No Spacing'
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 0)
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Tablas Markdown
        if '|' in line and not line.strip().startswith('#'):
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
            i += 1
            # Verificar si la siguiente línea no es una tabla
            if i >= len(lines) or '|' not in lines[i]:
                # Procesar la tabla
                process_table(doc, table_lines)
                in_table = False
                table_lines = []
            continue
        
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
        
        # Listas
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            # Procesar checkboxes
            text = text.replace('✅', '☑')
            text = text.replace('⚠️', '⚠')
            text = text.replace('❌', '☐')
            p = doc.add_paragraph(text, style='List Bullet')
        
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            p = doc.add_paragraph(text, style='List Number')
        
        # Separadores
        elif line.strip() == '---':
            p = doc.add_paragraph()
            p.add_run('_' * 80)
        
        # Blockquotes
        elif line.startswith('> '):
            text = line[2:]
            # Eliminar tags de tipo de alerta
            text = re.sub(r'\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]', '', text)
            if text.strip():
                p = doc.add_paragraph(text)
                p.style = 'Quote'
        
        # Texto normal
        elif line.strip():
            # Procesar negrita, cursiva, código inline
            text = process_inline_formatting(line)
            if text.strip():
                p = doc.add_paragraph(text)
        
        # Línea vacía
        else:
            if i > 0 and lines[i-1].strip():  # Evitar múltiples líneas vacías
                doc.add_paragraph()
        
        i += 1
    
    # Guardar documento
    doc.save(docx_file)
    print(f"✅ Documento Word creado exitosamente: {docx_file}")

def process_table(doc, table_lines):
    """Procesa una tabla Markdown y la añade al documento"""
    if len(table_lines) < 2:
        return
    
    # Filtrar líneas de separadores
    data_lines = [line for line in table_lines if not re.match(r'^\|[\s\-:|]+\|$', line)]
    
    if len(data_lines) < 1:
        return
    
    # Parsear filas
    rows = []
    for line in data_lines:
        cells = [cell.strip() for cell in line.split('|')]
        cells = [c for c in cells if c]  # Eliminar celdas vacías de los bordes
        if cells:
            rows.append(cells)
    
    if not rows:
        return
    
    # Crear tabla en Word
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    # Llenar la tabla
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.rows[i].cells[j]
                cell.text = process_inline_formatting(cell_text)
                # Primera fila como encabezado
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

def process_inline_formatting(text):
    """Procesa formato inline de Markdown (negrita, cursiva, código)"""
    # Eliminar código inline
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # Eliminar negrita
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    # Eliminar cursiva
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    # Eliminar links
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    return text

if __name__ == '__main__':
    md_path = Path(r'C:\Users\samue\.gemini\antigravity\brain\8913602d-3499-479c-b793-5862292333e2\estructura_proyecto.md')
    docx_path = Path(r'c:\Users\samue\Desktop\farmacia-saas\Reporte_Estructura_Proyecto.docx')
    
    print(f"📄 Convirtiendo {md_path.name} a Word...")
    convert_md_to_docx(md_path, docx_path)
